#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gerador de Relatório do Projeto Final - Módulo de Compras
Formato: Word (.docx) com normas ABNT
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def configurar_estilos_abnt(doc):
    """Configura estilos conforme normas ABNT"""
    
    # Estilo Normal - Texto do corpo
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_after = Pt(6)
    paragraph_format.first_line_indent = Inches(0.5)
    
    # Estilo para Títulos
    for i in range(1, 4):
        heading = doc.styles[f'Heading {i}']
        heading.font.name = 'Arial'
        heading.font.size = Pt(12)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

def adicionar_capa(doc):
    """Adiciona capa do trabalho conforme ABNT"""
    
    # Instituição
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('UNIVERSIDADE FEDERAL\n')
    run.font.size = Pt(12)
    run.font.bold = True
    
    run = p.add_run('CURSO DE CIÊNCIA DA COMPUTAÇÃO\n')
    run.font.size = Pt(12)
    run.font.bold = True
    
    run = p.add_run('DISCIPLINA: PROGRAMAÇÃO ORIENTADA A OBJETOS')
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Espaçamento
    for _ in range(8):
        doc.add_paragraph()
    
    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SISTEMA ERP - MÓDULO DE COMPRAS\n')
    run.font.size = Pt(14)
    run.font.bold = True
    
    run = p.add_run('Projeto Final de Programação Orientada a Objetos')
    run.font.size = Pt(12)
    
    # Espaçamento
    for _ in range(10):
        doc.add_paragraph()
    
    # Autor e Data
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n\n\nBrasília - DF\n{datetime.now().strftime("%B de %Y").upper()}')
    run.font.size = Pt(12)
    
    doc.add_page_break()

def adicionar_sumario(doc):
    """Adiciona sumário"""
    doc.add_heading('SUMÁRIO', 0)
    
    sumario_items = [
        ('1. INTRODUÇÃO', '3'),
        ('2. OBJETIVOS', '4'),
        ('   2.1 Objetivo Geral', '4'),
        ('   2.2 Objetivos Específicos', '4'),
        ('3. FUNDAMENTAÇÃO TEÓRICA', '5'),
        ('   3.1 Programação Orientada a Objetos', '5'),
        ('   3.2 Padrões de Projeto', '5'),
        ('4. ARQUITETURA DO SISTEMA', '6'),
        ('   4.1 Visão Geral', '6'),
        ('   4.2 Módulos do Sistema', '7'),
        ('   4.3 Diagrama de Classes', '8'),
        ('5. IMPLEMENTAÇÃO', '9'),
        ('   5.1 Tecnologias Utilizadas', '9'),
        ('   5.2 Estrutura de Diretórios', '9'),
        ('   5.3 Classes Principais', '10'),
        ('6. FUNCIONALIDADES', '12'),
        ('   6.1 Módulo de Fornecedores', '12'),
        ('   6.2 Módulo de Ordens de Compra', '13'),
        ('   6.3 Integração com Estoque', '14'),
        ('   6.4 Integração com Produção', '15'),
        ('   6.5 Integração com Financeiro', '16'),
        ('7. CONCEITOS DE POO APLICADOS', '17'),
        ('   7.1 Encapsulamento', '17'),
        ('   7.2 Herança', '17'),
        ('   7.3 Polimorfismo', '18'),
        ('   7.4 Abstração', '18'),
        ('8. TESTES E VALIDAÇÃO', '19'),
        ('9. CONCLUSÃO', '20'),
        ('REFERÊNCIAS', '21'),
    ]
    
    for item, pagina in sumario_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0) if not item.startswith('   ') else Inches(0.5)
        p.add_run(f'{item}').font.size = Pt(12)
        p.add_run('\t' * 5)
        p.add_run(pagina).font.size = Pt(12)
    
    doc.add_page_break()

def adicionar_introducao(doc):
    """Adiciona seção de introdução"""
    doc.add_heading('1. INTRODUÇÃO', 1)
    
    texto = """O presente trabalho apresenta o desenvolvimento de um Sistema de Gestão Empresarial (ERP), com foco no Módulo de Compras, desenvolvido como projeto final da disciplina de Programação Orientada a Objetos. O sistema foi implementado utilizando a linguagem C++ e aplicando os principais conceitos e paradigmas da programação orientada a objetos.

O Módulo de Compras é um componente crítico de qualquer sistema ERP, sendo responsável pela gestão de fornecedores, criação e acompanhamento de ordens de compra, além da integração com outros módulos essenciais como Estoque, Produção e Financeiro. A escolha deste módulo como tema do projeto justifica-se pela sua complexidade e relevância prática em ambientes corporativos.

O desenvolvimento do sistema seguiu as melhores práticas de engenharia de software, incluindo o uso de padrões de projeto, programação concorrente com threads e mutex, persistência de dados, tratamento de exceções, e uma arquitetura modular que facilita a manutenção e expansão do código.

Este relatório descreve detalhadamente a arquitetura, implementação e funcionalidades do sistema desenvolvido, demonstrando a aplicação prática dos conceitos de Programação Orientada a Objetos estudados durante o curso."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_objetivos(doc):
    """Adiciona seção de objetivos"""
    doc.add_heading('2. OBJETIVOS', 1)
    
    doc.add_heading('2.1 Objetivo Geral', 2)
    p = doc.add_paragraph('Desenvolver um Módulo de Compras completo e funcional para um sistema ERP, aplicando os conceitos fundamentais de Programação Orientada a Objetos, demonstrando compreensão prática dos paradigmas de encapsulamento, herança, polimorfismo e abstração.')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('2.2 Objetivos Específicos', 2)
    
    objetivos = [
        'Implementar um sistema de gerenciamento de fornecedores com validação de CNPJ e ordenação por diferentes critérios',
        'Desenvolver um sistema de ordens de compra com controle de status e integração com múltiplos módulos',
        'Aplicar programação concorrente utilizando threads e mutex para simulação de processos assíncronos',
        'Implementar persistência de dados em arquivos texto para garantir a durabilidade das informações',
        'Criar interfaces para integração com os módulos de Estoque, Produção e Financeiro',
        'Desenvolver uma interface de usuário em modo console intuitiva e completa',
        'Aplicar tratamento de exceções para garantir a robustez do sistema',
        'Utilizar estruturas de dados genéricas com templates para reutilização de código',
        'Implementar o padrão de projeto Mock Object para simulação de módulos externos',
        'Documentar o código seguindo as melhores práticas de engenharia de software'
    ]
    
    for obj in objetivos:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_fundamentacao(doc):
    """Adiciona fundamentação teórica"""
    doc.add_heading('3. FUNDAMENTAÇÃO TEÓRICA', 1)
    
    doc.add_heading('3.1 Programação Orientada a Objetos', 2)
    texto = """A Programação Orientada a Objetos (POO) é um paradigma de programação baseado no conceito de "objetos", que podem conter dados na forma de atributos e código na forma de métodos. Os quatro pilares fundamentais da POO são:

Encapsulamento: Mecanismo que restringe o acesso direto aos componentes de um objeto, protegendo o estado interno e expondo apenas interfaces públicas bem definidas.

Herança: Permite que classes derivadas herdem atributos e métodos de classes base, promovendo a reutilização de código e estabelecendo relações hierárquicas.

Polimorfismo: Capacidade de objetos de diferentes classes responderem de forma diferente à mesma mensagem ou chamada de método, através de sobrescrita ou sobrecarga.

Abstração: Processo de simplificação da realidade, focando nos aspectos essenciais e escondendo detalhes de implementação através de classes abstratas e interfaces."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('3.2 Padrões de Projeto', 2)
    texto = """O projeto implementa diversos padrões de projeto reconhecidos pela engenharia de software:

Mock Object Pattern: Utilizado para simular o comportamento de módulos externos (Estoque, Produção, Financeiro) sem implementar toda a complexidade real desses sistemas.

Template Method: Aplicado na estrutura genérica ListaGenerica<T>, permitindo operações em qualquer tipo de dado.

Facade Pattern: A classe ModuloCompras atua como uma fachada, simplificando o acesso aos subsistemas internos (GerenciadorFornecedores, GerenciadorOrdens, PersistenciaCompras)."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_arquitetura(doc):
    """Adiciona seção de arquitetura"""
    doc.add_heading('4. ARQUITETURA DO SISTEMA', 1)
    
    doc.add_heading('4.1 Visão Geral', 2)
    texto = """O sistema foi projetado seguindo uma arquitetura em camadas, com clara separação de responsabilidades entre os componentes. A arquitetura é composta por três camadas principais:

Camada de Apresentação: Implementada no arquivo main.cpp, responsável pela interface com o usuário através de um menu interativo em modo console.

Camada de Lógica de Negócio: Contém as classes de gerenciamento (ModuloCompras, GerenciadorFornecedores, GerenciadorOrdens) que implementam as regras de negócio do sistema.

Camada de Dados: Implementada pela classe PersistenciaCompras, responsável pela leitura e escrita de dados em arquivos.

A comunicação entre as camadas é feita através de interfaces bem definidas, garantindo baixo acoplamento e alta coesão."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('4.2 Módulos do Sistema', 2)
    
    modulos = {
        'Módulo de Fornecedores': 'Gerencia o cadastro, listagem, busca e remoção de fornecedores. Inclui validação de CNPJ e funcionalidades de ordenação por preço.',
        'Módulo de Ordens de Compra': 'Controla o ciclo de vida das ordens de compra, desde a criação até a aprovação/rejeição, incluindo integração com módulos externos.',
        'Módulo de Persistência': 'Responsável pela serialização e deserialização de dados em arquivos texto, garantindo a durabilidade das informações.',
        'Módulo de Integração - Estoque': 'Interface para comunicação com o sistema de estoque, registrando entradas de materiais e consultando disponibilidade.',
        'Módulo de Integração - Produção': 'Interface para comunicação com o sistema de produção, recebendo pedidos de materiais e notificando sobre compras realizadas.',
        'Módulo de Integração - Financeiro': 'Interface para comunicação com o sistema financeiro, verificando disponibilidade de verba e registrando contas a pagar.'
    }
    
    for nome, descricao in modulos.items():
        p = doc.add_paragraph()
        p.add_run(f'{nome}: ').bold = True
        p.add_run(descricao)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_diagrama_classes(doc):
    """Adiciona explicação do diagrama de classes"""
    doc.add_heading('4.3 Diagrama de Classes', 2)
    
    texto = """O diagrama de classes do sistema ilustra a estrutura completa do Módulo de Compras, mostrando as principais classes, suas relações e dependências. As classes principais são:

Classes de Domínio:
- Pessoa: Classe base abstrata que define atributos comuns (nome, endereço)
- Fornecedor: Especialização de Pessoa, adiciona CNPJ, produto e preço
- OrdemCompra: Representa uma ordem de compra com seus atributos e status

Classes de Gerenciamento:
- ModuloCompras: Classe coordenadora principal (Facade)
- GerenciadorFornecedores: Gerencia a coleção de fornecedores
- GerenciadorOrdens: Gerencia ordens de compra e integrações

Classes de Infraestrutura:
- ListaGenerica<T>: Template para estrutura de dados genérica
- PersistenciaCompras: Gerencia persistência em arquivos
- ComprasException: Classe de exceção customizada

Interfaces de Integração:
- IEstoque: Interface para módulo de estoque
- IProducao: Interface para módulo de produção
- IFinanceiro: Interface para módulo financeiro

Classes Mock:
- EstoqueMock: Implementação simulada do estoque
- ProducaoMock: Implementação simulada da produção
- FinanceiroMock: Implementação simulada do financeiro

O diagrama completo pode ser visualizado no arquivo 'diagrama de classes.pdf' anexo a este relatório."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_implementacao(doc):
    """Adiciona seção de implementação"""
    doc.add_heading('5. IMPLEMENTAÇÃO', 1)
    
    doc.add_heading('5.1 Tecnologias Utilizadas', 2)
    
    tecnologias = [
        'Linguagem: C++ (Padrão C++17)',
        'Compilador: g++ (GNU Compiler Collection)',
        'Paradigma: Programação Orientada a Objetos',
        'Bibliotecas Padrão: iostream, fstream, string, vector, map, thread, mutex, chrono',
        'Controle de Versão: Git',
        'Sistema Operacional: Linux (Ubuntu 24.04)',
        'IDE/Editor: Visual Studio Code'
    ]
    
    for tec in tecnologias:
        p = doc.add_paragraph(tec, style='List Bullet')
    
    doc.add_heading('5.2 Estrutura de Diretórios', 2)
    
    p = doc.add_paragraph('A organização do projeto segue uma estrutura modular e bem definida:')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    estrutura = """
MODULO DE COMPRAS/
│
├── include/              # Arquivos de cabeçalho (.h)
│   ├── ComprasException.h
│   ├── EstoqueMock.h
│   ├── FinanceiroMock.h
│   ├── Fornecedor.h
│   ├── GerenciadorFornecedores.h
│   ├── GerenciadorOrdens.h
│   ├── IEstoque.h
│   ├── IExibivel.h
│   ├── IFinanceiro.h
│   ├── IProducao.h
│   ├── ListaGenerica.h
│   ├── ModuloCompras.h
│   ├── OrdemCompra.h
│   ├── PersistenciaCompras.h
│   ├── Pessoa.h
│   └── ProducaoMock.h
│
├── src/                  # Arquivos de implementação (.cpp)
│   ├── GerenciadorFornecedores.cpp
│   ├── GerenciadorOrdens.cpp
│   ├── main.cpp
│   ├── ModuloCompras.cpp
│   └── PersistenciaCompras.cpp
│
├── bin/                  # Executável compilado
│   └── modulo_compras
│
├── data/                 # Arquivos de dados persistentes
│   ├── fornecedores.txt
│   └── ordens.txt
│
└── docs/                 # Documentação
    ├── TESTES.md
    └── README_STATUS.md
"""
    
    p = doc.add_paragraph(estrutura)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_page_break()

def adicionar_classes_principais(doc):
    """Adiciona descrição das classes principais"""
    doc.add_heading('5.3 Classes Principais', 2)
    
    classes = {
        'Pessoa (Classe Abstrata)': """Classe base que implementa o conceito de abstração da POO. Define atributos comuns (nome, endereço) e o método virtual puro exibirDetalhes(), forçando classes derivadas a implementarem sua própria versão.""",
        
        'Fornecedor': """Herda de Pessoa e adiciona atributos específicos: CNPJ, produto e preço. Implementa validação de CNPJ através de algoritmo de validação de dígitos verificadores. Demonstra o conceito de herança e especialização.""",
        
        'OrdemCompra': """Representa uma ordem de compra no sistema. Contém atributos como ID, item, quantidade, valor unitário, fornecedor e status (PENDENTE, APROVADO, REJEITADO). Implementa encapsulamento através de getters e setters.""",
        
        'GerenciadorFornecedores': """Gerencia a coleção de fornecedores utilizando ListaGenerica<Fornecedor>. Implementa operações CRUD (Create, Read, Update, Delete) e funcionalidades especiais como ordenação por preço e busca por produto.""",
        
        'GerenciadorOrdens': """Classe mais complexa do sistema, responsável por criar e gerenciar ordens de compra. Implementa:
        - Programação concorrente com threads para verificação de verba
        - Mutex para proteção de recursos compartilhados
        - Integração com módulos externos (Estoque, Produção, Financeiro)
        - Cálculo de estatísticas""",
        
        'ModuloCompras (Facade)': """Classe coordenadora que atua como fachada para o sistema. Simplifica o acesso aos subsistemas (GerenciadorFornecedores, GerenciadorOrdens, PersistenciaCompras) através de uma interface unificada.""",
        
        'PersistenciaCompras': """Responsável pela persistência de dados em arquivos texto. Implementa serialização e deserialização de fornecedores e ordens de compra, garantindo a durabilidade dos dados entre execuções.""",
        
        'ListaGenerica<T> (Template)': """Estrutura de dados genérica implementada com template. Demonstra o conceito de genericidade e reutilização de código. Permite armazenar qualquer tipo de objeto.""",
        
        'ComprasException': """Classe de exceção customizada derivada de std::exception. Utilizada para tratamento de erros específicos do domínio de compras."""
    }
    
    for classe, descricao in classes.items():
        p = doc.add_paragraph()
        p.add_run(f'\n{classe}\n').bold = True
        p.add_run(descricao)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_funcionalidades(doc):
    """Adiciona seção de funcionalidades"""
    doc.add_heading('6. FUNCIONALIDADES', 1)
    
    doc.add_heading('6.1 Módulo de Fornecedores', 2)
    
    funcionalidades_fornecedores = {
        'Cadastrar Fornecedor': 'Permite adicionar novos fornecedores ao sistema com validação de CNPJ. Valida duplicidade e formato correto do CNPJ através de algoritmo de dígitos verificadores.',
        
        'Listar Fornecedores': 'Exibe todos os fornecedores cadastrados com suas informações completas: ID, nome, endereço, CNPJ, produto e preço.',
        
        'Buscar Fornecedor por ID': 'Localiza um fornecedor específico através do seu identificador único.',
        
        'Buscar Fornecedores por Produto': 'Lista todos os fornecedores que oferecem um determinado produto.',
        
        'Ordenar por Preço': 'Ordena a lista de fornecedores por preço do produto, auxiliando na escolha do fornecedor mais econômico.',
        
        'Remover Fornecedor': 'Permite excluir um fornecedor do sistema (não implementado na interface atual por segurança).'
    }
    
    for funcao, descricao in funcionalidades_fornecedores.items():
        p = doc.add_paragraph()
        p.add_run(f'{funcao}: ').bold = True
        p.add_run(descricao)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('6.2 Módulo de Ordens de Compra', 2)
    
    funcionalidades_ordens = {
        'Criar Ordem de Compra': """Funcionalidade principal que demonstra a integração completa do sistema. O processo de criação inclui:
        1. Validação de entrada (quantidade e valor positivos)
        2. Proteção com mutex para operações thread-safe
        3. Criação da ordem com status PENDENTE
        4. Verificação de verba disponível (thread paralela com latência simulada)
        5. Autorização de pagamento pelo financeiro
        6. Registro como conta a pagar no financeiro
        7. Notificação ao módulo de produção
        8. Atualização de previsão de entrega
        9. Registro de entrada no estoque
        10. Atualização do status para APROVADO ou REJEITADO""",
        
        'Listar Ordens de Compra': 'Exibe todas as ordens cadastradas com status, valores e detalhes completos.',
        
        'Buscar Ordem por ID': 'Localiza uma ordem específica através do seu identificador.',
        
        'Exibir Estatísticas': 'Calcula e exibe estatísticas das ordens: total de aprovadas, rejeitadas, pendentes e valor total aprovado.'
    }
    
    for funcao, descricao in funcionalidades_ordens.items():
        p = doc.add_paragraph()
        p.add_run(f'\n{funcao}:\n').bold = True
        p.add_run(descricao)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_integracoes(doc):
    """Adiciona detalhes das integrações"""
    doc.add_heading('6.3 Integração com Estoque', 2)
    
    texto = """O módulo de compras integra-se com o estoque através da interface IEstoque, que define os seguintes métodos:

consultarItem(idMaterial): Consulta a disponibilidade de um material específico no estoque.

listarTodosItens(): Lista todos os itens presentes no estoque com suas quantidades.

registrarEntradaCompra(idMaterial, quantidade, idOrdemCompra): Registra a entrada de material no estoque quando uma compra é aprovada. Este método é chamado automaticamente durante o processo de criação de ordem de compra.

reservarMaterial(idMaterial, quantidade): Permite reservar material do estoque para produção ou outras finalidades.

A implementação EstoqueMock simula um sistema de estoque real, mantendo um inventário em memória com três materiais pré-cadastrados: Aço Inox (100 unidades), Parafusos M10 (500 unidades) e Borracha Industrial (50 unidades)."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('6.4 Integração com Produção', 2)
    
    texto = """O módulo de produção é integrado através da interface IProducao, que define:

notificarMaterialComprado(idMaterial): Notifica o módulo de produção quando um material é comprado com sucesso.

receberPedidoMaterial(idMaterial, quantidade, prioridade): Permite que a produção solicite materiais ao módulo de compras, especificando prioridade (1-Baixa, 2-Média, 3-Alta).

atualizarPrevisaoEntrega(idOrdemCompra, dataPrevisao): Informa à produção a previsão de entrega de um material comprado.

listarPedidosPendentes(): Lista todos os pedidos de materiais que ainda não foram atendidos.

A implementação ProducaoMock mantém uma lista de pedidos com dois pedidos pré-cadastrados: Material 1 (50 unidades, prioridade alta) e Material 2 (200 unidades, prioridade média)."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('6.5 Integração com Financeiro', 2)
    
    texto = """O módulo financeiro é essencial para o controle de aprovação de compras. A interface IFinanceiro define:

verificarDisponibilidade(valor): Verifica se há verba disponível para realizar uma compra. Este método é executado em uma thread separada, simulando latência de comunicação (2-4 segundos).

autorizarPagamento(idOrdem): Autoriza efetivamente o pagamento de uma ordem de compra após verificação de disponibilidade.

registrarContaPagar(idOrdemCompra, valorTotal, fornecedor, dataVencimento): Registra a compra como uma conta a pagar no sistema financeiro, incluindo informações do fornecedor e data de vencimento.

listarContasPagar(): Lista todas as contas a pagar registradas no sistema.

A implementação FinanceiroMock simula um orçamento de R$ 100.000,00 e mantém registro de todas as contas a pagar geradas pelas ordens de compra aprovadas."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_conceitos_poo(doc):
    """Adiciona seção sobre conceitos de POO"""
    doc.add_heading('7. CONCEITOS DE POO APLICADOS', 1)
    
    doc.add_heading('7.1 Encapsulamento', 2)
    texto = """O encapsulamento foi amplamente aplicado em todas as classes do sistema:

Atributos Privados: Todos os atributos das classes são declarados como private, protegendo o estado interno dos objetos.

Métodos Públicos: O acesso aos atributos é feito exclusivamente através de métodos getter e setter públicos, permitindo validação e controle.

Exemplo na classe Fornecedor:
- Atributos private: cnpj, produto, precoProduto
- Métodos public: getCNPJ(), getProduto(), setPrecoProduto()
- Validação no construtor para garantir CNPJ válido

Benefícios: O encapsulamento garantiu que os dados sejam manipulados apenas de forma controlada, evitando estados inválidos e facilitando manutenção."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('7.2 Herança', 2)
    texto = """A herança foi implementada na hierarquia de classes de domínio:

Classe Base Pessoa:
- Define atributos comuns: nome, endereço
- Método virtual puro: exibirDetalhes()
- Construtor protegido

Classe Derivada Fornecedor:
- Herda atributos de Pessoa
- Adiciona atributos específicos: CNPJ, produto, preço
- Implementa exibirDetalhes() com informações específicas
- Demonstra especialização e reutilização de código

Esta hierarquia demonstra o relacionamento "é-um" (Fornecedor é uma Pessoa) e permite tratamento polimórfico através da classe base."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('7.3 Polimorfismo', 2)
    texto = """O polimorfismo foi aplicado através de múltiplas técnicas:

Polimorfismo de Subtipo:
- Interface IExibivel com método virtual puro exibir()
- Implementações específicas em Fornecedor e OrdemCompra
- Permite tratar objetos diferentes através da mesma interface

Polimorfismo Paramétrico:
- Template ListaGenerica<T> aceita qualquer tipo
- Mesmos métodos funcionam com Fornecedor e OrdemCompra
- Demonstra genericidade e reutilização

Sobrecarga de Operadores:
- Operador de inserção (<<) sobrecarregado para OrdemCompra
- Facilita a saída formatada de objetos

O polimorfismo permitiu código mais flexível e extensível, facilitando a adição de novos tipos sem modificar código existente."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading('7.4 Abstração', 2)
    texto = """A abstração foi fundamental para simplificar a complexidade do sistema:

Classes Abstratas:
- Pessoa define interface comum sem implementação completa
- Força classes derivadas a implementarem comportamentos específicos

Interfaces:
- IEstoque, IProducao, IFinanceiro definem contratos claros
- Ocultam detalhes de implementação dos módulos externos
- Permitem substituir implementações sem afetar o sistema

Padrão Facade:
- ModuloCompras abstrai a complexidade dos subsistemas
- Interface simplificada para operações complexas
- Usuário não precisa conhecer detalhes internos

A abstração permitiu focar no "o que" fazer ao invés do "como" fazer, resultando em código mais limpo e manutenível."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_testes(doc):
    """Adiciona seção de testes"""
    doc.add_heading('8. TESTES E VALIDAÇÃO', 1)
    
    texto = """O sistema foi extensivamente testado para garantir sua funcionalidade e robustez. Os testes realizados incluem:

Testes Unitários:
- Validação de CNPJ com casos válidos e inválidos
- Operações da ListaGenerica (adicionar, remover, buscar)
- Cálculo de estatísticas de ordens de compra
- Serialização e deserialização de dados

Testes de Integração:
- Fluxo completo de criação de ordem de compra
- Integração com módulos Estoque, Produção e Financeiro
- Persistência e recuperação de dados
- Comunicação entre threads

Testes de Concorrência:
- Múltiplas threads acessando recursos compartilhados
- Proteção com mutex funcionando corretamente
- Ausência de race conditions e deadlocks

Testes de Interface:
- Todas as opções do menu funcionando corretamente
- Validação de entrada do usuário
- Tratamento de erros e exceções

Testes de Performance:
- Tempo de resposta do sistema
- Latência simulada do módulo financeiro
- Capacidade de gerenciar múltiplos fornecedores e ordens

Resultados:
- ✓ Compilação sem erros ou warnings
- ✓ Todas as funcionalidades operacionais
- ✓ Integrações funcionando conforme especificado
- ✓ Persistência de dados validada
- ✓ Sistema robusto e estável

Arquivos de Teste:
Os scripts de teste estão disponíveis no diretório do projeto:
- teste_modulos.sh: Testa módulos individuais
- TESTES.md: Documentação completa dos testes
- README_STATUS.md: Status atual do sistema"""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_conclusao(doc):
    """Adiciona conclusão"""
    doc.add_heading('9. CONCLUSÃO', 1)
    
    texto = """O desenvolvimento do Sistema ERP - Módulo de Compras proporcionou uma aplicação prática abrangente dos conceitos de Programação Orientada a Objetos estudados ao longo do curso. O projeto demonstrou com sucesso a implementação dos quatro pilares fundamentais da POO: encapsulamento, herança, polimorfismo e abstração.

A arquitetura modular adotada, com clara separação de responsabilidades e uso de padrões de projeto reconhecidos, resultou em um sistema robusto, manutenível e extensível. A aplicação de conceitos avançados como programação concorrente com threads e mutex, tratamento de exceções, templates e interfaces demonstra uma compreensão profunda dos paradigmas de programação orientada a objetos.

As integrações implementadas com os módulos de Estoque, Produção e Financeiro ilustram cenários reais de sistemas corporativos, onde diferentes subsistemas precisam comunicar-se de forma eficiente e confiável. A utilização do padrão Mock Object permitiu simular essas integrações sem a necessidade de implementar sistemas completos, demonstrando flexibilidade e pragmatismo no desenvolvimento.

Os testes realizados validaram a funcionalidade e robustez do sistema, comprovando que todos os requisitos especificados foram atendidos. O código foi desenvolvido seguindo as melhores práticas de engenharia de software, incluindo documentação adequada, nomenclatura clara e organização lógica dos arquivos.

Como trabalhos futuros, sugere-se:
- Implementação de uma interface gráfica (GUI)
- Expansão para outros módulos do ERP (Vendas, RH, etc.)
- Implementação de banco de dados relacional
- Adição de relatórios e dashboard analítico
- Implementação de testes automatizados com frameworks de teste

O projeto alcançou plenamente seus objetivos, demonstrando aplicação prática e compreensão teórica dos conceitos de Programação Orientada a Objetos, constituindo uma base sólida para o desenvolvimento de sistemas complexos e de qualidade profissional."""
    
    for paragrafo in texto.split('\n\n'):
        p = doc.add_paragraph(paragrafo)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

def adicionar_referencias(doc):
    """Adiciona referências bibliográficas"""
    doc.add_heading('REFERÊNCIAS', 1)
    
    referencias = [
        'DEITEL, P.; DEITEL, H. C++: Como Programar. 10. ed. São Paulo: Pearson, 2017.',
        
        'GAMMA, E. et al. Padrões de Projeto: Soluções Reutilizáveis de Software Orientado a Objetos. Porto Alegre: Bookman, 2000.',
        
        'MARTIN, R. C. Código Limpo: Habilidades Práticas do Agile Software. Rio de Janeiro: Alta Books, 2009.',
        
        'STROUSTRUP, B. The C++ Programming Language. 4. ed. Upper Saddle River: Addison-Wesley, 2013.',
        
        'SOMMERVILLE, I. Engenharia de Software. 10. ed. São Paulo: Pearson, 2018.',
        
        'PRESSMAN, R. S. Engenharia de Software: Uma Abordagem Profissional. 8. ed. Porto Alegre: AMGH, 2016.',
        
        'ISO/IEC 14882:2017. Programming Languages — C++. International Organization for Standardization, 2017.',
        
        'CPPREFERENCE. C++ Reference. Disponível em: <https://en.cppreference.com/>. Acesso em: 25 nov. 2025.'
    ]
    
    for ref in referencias:
        p = doc.add_paragraph(ref)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

def gerar_relatorio():
    """Função principal que gera o relatório completo"""
    print("Gerando Relatório do Projeto...")
    
    # Criar documento
    doc = Document()
    
    # Configurar estilos ABNT
    configurar_estilos_abnt(doc)
    
    # Adicionar seções
    print("  - Adicionando capa...")
    adicionar_capa(doc)
    
    print("  - Adicionando sumário...")
    adicionar_sumario(doc)
    
    print("  - Adicionando introdução...")
    adicionar_introducao(doc)
    
    print("  - Adicionando objetivos...")
    adicionar_objetivos(doc)
    
    print("  - Adicionando fundamentação teórica...")
    adicionar_fundamentacao(doc)
    
    print("  - Adicionando arquitetura...")
    adicionar_arquitetura(doc)
    
    print("  - Adicionando diagrama de classes...")
    adicionar_diagrama_classes(doc)
    
    print("  - Adicionando implementação...")
    adicionar_implementacao(doc)
    
    print("  - Adicionando classes principais...")
    adicionar_classes_principais(doc)
    
    print("  - Adicionando funcionalidades...")
    adicionar_funcionalidades(doc)
    
    print("  - Adicionando integrações...")
    adicionar_integracoes(doc)
    
    print("  - Adicionando conceitos de POO...")
    adicionar_conceitos_poo(doc)
    
    print("  - Adicionando testes...")
    adicionar_testes(doc)
    
    print("  - Adicionando conclusão...")
    adicionar_conclusao(doc)
    
    print("  - Adicionando referências...")
    adicionar_referencias(doc)
    
    # Salvar documento
    filename = 'Relatorio_Projeto_Final_Modulo_Compras.docx'
    doc.save(filename)
    
    print(f"\n✓ Relatório gerado com sucesso: {filename}")
    print(f"  Localização: /workspaces/POO-PROJETO-FINAL/MODULO DE COMPRAS/{filename}")
    print("\nO documento contém:")
    print("  - Capa formatada conforme ABNT")
    print("  - Sumário completo")
    print("  - 9 capítulos principais")
    print("  - Descrição detalhada de todas as funcionalidades")
    print("  - Explicação dos conceitos de POO aplicados")
    print("  - Referências bibliográficas")
    print("\nNOTA: Para incluir o diagrama de classes visual,")
    print("      abra o arquivo Word e insira manualmente a imagem")
    print("      do arquivo 'diagrama de classes.pdf' na seção 4.3")

if __name__ == '__main__':
    gerar_relatorio()
